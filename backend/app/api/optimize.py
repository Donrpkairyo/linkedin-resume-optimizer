from fastapi import APIRouter, HTTPException, File, UploadFile, Form, Response
from typing import Optional
from datetime import datetime
import io
import os
import docx
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor

from app.models.base import ResumeOptimizationRequest, ResumeOptimizationResponse
from app.services.gemini_service import GeminiService
from app.services.linkedin_service import LinkedInService

router = APIRouter()
gemini_service = GeminiService()
linkedin_service = LinkedInService()

@router.post("/resume", response_model=ResumeOptimizationResponse)
async def optimize_resume(request: ResumeOptimizationRequest) -> ResumeOptimizationResponse:
    """
    Optimize resume based on a job description.
    """
    try:
        suggestions = await gemini_service.optimize_resume(request.job_description, request.resume_text)
        return ResumeOptimizationResponse(
            original_resume=request.resume_text,
            optimized_resume=suggestions,
            optimization_id=str(hash(request.resume_text + request.job_description)),
            created_at=datetime.now().isoformat()
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Resume optimization failed: {str(e)}"
        )

@router.post("/resume/url", response_model=ResumeOptimizationResponse)
async def optimize_resume_from_url(resume_text: str, job_url: str) -> ResumeOptimizationResponse:
    """
    Optimize resume based on a LinkedIn job URL.
    """
    try:
        # Get job ID from URL
        job_id = job_url.split('/')[-1].split('?')[0]
        
        # Fetch job description
        description = await linkedin_service.get_job_description(job_id)
        if not description:
            raise HTTPException(
                status_code=404,
                detail="Could not fetch job description"
            )

        # Generate optimization suggestions
        suggestions = await gemini_service.optimize_resume(description, resume_text)
        
        return ResumeOptimizationResponse(
            original_resume=resume_text,
            optimized_resume=suggestions,
            optimization_id=str(hash(resume_text + description)),
            created_at=datetime.now().isoformat()
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Resume optimization failed: {str(e)}"
        )

@router.post("/resume/docx", response_model=ResumeOptimizationResponse)
async def optimize_resume_from_docx(
    resume: UploadFile = File(...),
    job_description: str = Form(""),
    job_url: str = Form("")
) -> ResumeOptimizationResponse:
    """
    Optimize resume from a DOCX file based on a job description or URL.
    """
    try:
        # Validate resume file
        if not resume.filename:
            raise HTTPException(
                status_code=400,
                detail="No file provided"
            )

        if not resume.filename.endswith(('.docx', '.doc')):
            raise HTTPException(
                status_code=400,
                detail="Only .doc and .docx files are supported"
            )

        # Read resume content
        try:
            content = await resume.read()
            if not content:
                raise ValueError("Empty file")
            
            doc = docx.Document(io.BytesIO(content))
            resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            if not resume_text:
                raise ValueError("No text content found in document")
                
        except ValueError as e:
            raise HTTPException(
                status_code=400,
                detail=str(e)
            )
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Error reading document: {str(e)}"
            )

        # Handle job details
        if job_url:
            try:
                # Basic URL validation
                if not job_url.startswith(('http://', 'https://')):
                    raise ValueError("Invalid URL format")
                    
                if 'linkedin.com/jobs' not in job_url:
                    raise ValueError("Not a LinkedIn job URL")
                
                # Extract and validate job ID
                job_id = job_url.split('/')[-1].split('?')[0]
                if not job_id.strip() or not job_id.isdigit():
                    raise ValueError("Invalid job ID in URL")
                
                # Fetch job description
                job_description = await linkedin_service.get_job_description(job_id)
                if not job_description:
                    raise ValueError("Could not fetch job description")
                    
            except ValueError as e:
                raise HTTPException(
                    status_code=400,
                    detail=str(e)
                )
        
        # Validate we have either job description or URL
        if not job_description and not job_url:
            raise HTTPException(
                status_code=400,
                detail="Either job description or job URL is required"
            )

        # Generate optimization suggestions
        try:
            suggestions = await gemini_service.optimize_resume(job_description, resume_text)
            if not suggestions:
                raise ValueError("Failed to generate optimization suggestions")
                
            return ResumeOptimizationResponse(
                original_resume=resume_text,
                optimized_resume=suggestions,
                optimization_id=str(hash(resume_text + job_description)),
                created_at=datetime.now().isoformat()
            )
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to optimize resume: {str(e)}"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Unexpected error: {str(e)}"
        )

@router.post("/resume/export")
async def export_resume(
    resume: UploadFile = File(...),
    suggestions: str = Form(...)
) -> Response:
    """
    Create a new resume file with the optimized content.
    """
    try:
        content = await resume.read()
        doc = docx.Document(io.BytesIO(content))
        
        # Parse position updates
        position_updates = {}
        current_position = None
        parsing_updates = False
        
        for line in suggestions.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            if line == "POSITION_UPDATES:":
                parsing_updates = True
                continue
            
            if not parsing_updates:
                continue
                
            if line.startswith('-'):
                if current_position:
                    position_updates.setdefault(current_position, []).append(line[1:].strip())
            else:
                current_position = line

        # Process document
        processed_positions = set()
        in_position = False
        current_position = None
        bullet_index = 0
        updates_made = False

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Check for position titles
            is_position = False
            if '|' in text and any(month in text for month in ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']):
                is_position = True

            # Process position titles
            if is_position:
                position_key = text.split('|')[0].strip().lower()
                for update_position in position_updates.keys():
                    if update_position.split('|')[0].strip().lower() == position_key and update_position not in processed_positions:
                        current_position = update_position
                        in_position = True
                        bullet_index = 0
                        processed_positions.add(update_position)
                        break

            # Update bullet points
            if in_position and bullet_index < len(position_updates.get(current_position, [])):
                if text.startswith('•') or any(word in text.lower() for word in ['increased', 'decreased', 'improved', 'achieved', 'launched', 'created', 'developed', 'implemented', 'managed', 'led']):
                    # Get original formatting
                    original_format = None
                    if paragraph.runs:
                        run = paragraph.runs[0]
                        original_format = {
                            'name': run.font.name or 'Arial',
                            'size': run.font.size,
                            'color': run.font.color.rgb if run.font.color else RGBColor(0, 0, 0)
                        }
                    
                    # Preserve bullet point
                    bullet = '• ' if text.startswith('•') else ''
                    
                    # Update content
                    new_content = position_updates[current_position][bullet_index]
                    paragraph.clear()
                    
                    # Add bullet if needed
                    if bullet:
                        run = paragraph.add_run(bullet)
                        if original_format:
                            run.font.name = original_format['name']
                            run.font.size = original_format['size']
                            if original_format['color']:
                                run.font.color.rgb = original_format['color']
                    
                    # Add content
                    run = paragraph.add_run(new_content.lstrip('• '))
                    if original_format:
                        run.font.name = original_format['name']
                        run.font.size = original_format['size']
                        if original_format['color']:
                            run.font.color.rgb = original_format['color']
                    
                    bullet_index += 1
                    updates_made = True

        # Save to memory
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return Response(
            content=output.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                'Content-Disposition': 'attachment; filename=optimized_resume.docx'
            }
        )
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to export resume: {str(e)}"
        )